import gradio as gr
import json
import re
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
import docx
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_COLOR_INDEX
import io
import tempfile
import os
from dotenv import load_dotenv
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import requests
from groq import Groq
import pickle
import zipfile
# PDF processing
import PyPDF2
from PyPDF2 import PdfReader

# Load environment variables
load_dotenv()

# Get Groq API key from environment variables
GROQ_API_KEY = os.getenv("groq_api_key")

class RAGSystem:
    def __init__(self):
        """Initialize RAG system with ADGM legal knowledge base for the Corporate Agent"""
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.index = None
        self.documents = []
        self.setup_adgm_knowledge_base()
    
    def setup_adgm_knowledge_base(self):
        """Initialize FAISS index with comprehensive ADGM legal documents"""
        # Comprehensive ADGM legal documents for RAG knowledge base
        adgm_documents = [
            {
                "source": "ADGM Companies Regulations 2020, Article 6",
                "content": "All companies incorporated in ADGM must specify ADGM Courts as the exclusive jurisdiction for dispute resolution. Any reference to UAE Federal Courts, Dubai Courts, or other jurisdictions shall render the clause invalid. The jurisdiction clause must explicitly state 'ADGM Courts' for legal validity."
            },
            {
                "source": "ADGM Companies Regulations 2020, Article 12",
                "content": "Articles of Association must contain the following mandatory provisions: (1) company objects clause, (2) authorized share capital structure with specific amounts, (3) director appointment and removal procedures, (4) shareholder meeting requirements including quorum and voting, (5) dividend distribution rights."
            },
            {
                "source": "ADGM Companies Regulations 2020, Article 25",
                "content": "Every company must maintain a register of members containing: full legal names, residential addresses, number of shares held, class of shares, date of entry as member, date of ceasing to be member. The register must be kept at the registered office and updated within 14 days of any change."
            },
            {
                "source": "ADGM Companies Regulations 2020, Article 18",
                "content": "Memorandum of Association must specify: (1) company name ending with 'Limited' or 'Ltd', (2) registered office address within ADGM, (3) company objects clause describing business activities, (4) authorized share capital amount in AED, (5) liability limitation clause."
            },
            {
                "source": "ADGM Employment Regulations 2019, Section 4.1",
                "content": "Employment contracts must specify in writing: job title and description, basic salary amount, working hours per week, annual leave entitlements, probation period, termination notice periods, and termination procedures. All contracts must be signed by both employer and employee with witness signatures."
            },
            {
                "source": "ADGM UBO Regulations 2021, Article 3",
                "content": "Ultimate Beneficial Owner declarations must identify all natural persons holding 25% or more ownership interest, directly or indirectly, including through trust arrangements. Declarations must include full name, nationality, passport details, residential address, percentage ownership, and nature of control. Updates required within 14 days of changes."
            },
            {
                "source": "ADGM Board Resolution Template Requirements",
                "content": "Board resolutions must include: (1) meeting date and time, (2) list of directors present and absent, (3) quorum confirmation statement, (4) resolved matters using specific format 'IT IS HEREBY RESOLVED THAT...', (5) signature blocks for all directors with printed names and dates, (6) company secretary attestation."
            },
            {
                "source": "ADGM MoA Template Standards 2023",
                "content": "Memorandum of Association template requirements: company name must end with 'Limited', registered office must be physical ADGM address, objects clause must be specific and comprehensive, share capital must be stated in AED with par value, subscriber details must include full names and addresses."
            },
            {
                "source": "ADGM Incorporation Checklist 2024",
                "content": "Mandatory documents for company incorporation: (1) Articles of Association - company constitution and rules, (2) Memorandum of Association - company formation document, (3) Board Resolution for incorporation - directors' authorization, (4) UBO Declaration Form - beneficial ownership disclosure, (5) Register of Members and Directors - shareholder and director details, (6) Incorporation Application Form - official registration request."
            },
            {
                "source": "ADGM Licensing Requirements 2024",
                "content": "Business licensing applications require: (1) License Application Form - official permit request, (2) Business Plan - detailed operational strategy, (3) Financial Projections - 3-year financial forecasts, (4) Professional Qualification Certificates - key personnel credentials, (5) Compliance Manual - operational procedures, (6) Board Resolution for licensing - director authorization."
            },
            {
                "source": "ADGM Employment Setup Requirements 2024",
                "content": "Employment framework establishment requires: (1) Employment Contract Template - standardized terms, (2) Employee Handbook - company policies, (3) HR Policies Manual - human resources procedures, (4) Workplace Health & Safety Policy - safety standards, (5) Data Protection Policy - privacy compliance measures."
            },
            {
                "source": "ADGM Contract Standards Guide",
                "content": "Legal documents must use definitive binding language creating clear obligations. Use 'shall', 'will', 'must' for mandatory provisions. Avoid ambiguous terms like 'may', 'might', 'could', 'possibly' in operative clauses as they create uncertainty and are non-binding. All operative clauses must be clear, specific, and enforceable."
            },
            {
                "source": "ADGM Signatory Requirements 2023",
                "content": "All legal documents require proper execution with: (1) authorized signatory signatures, (2) printed names below signatures, (3) signature dates, (4) witness signatures where required, (5) company seal affixation for corporate documents, (6) director capacity notation (e.g., 'Director'), (7) proper formatting with signature blocks."
            },
            {
                "source": "ADGM Shareholder Resolution Standards",
                "content": "Shareholder resolutions must contain: meeting notice period compliance, quorum requirements (minimum 2 shareholders or 25% of shares), resolution text using 'IT IS RESOLVED THAT' format, voting results with percentage approvals, chairman signature and date, company secretary certification."
            }
        ]
        
        self.documents = adgm_documents
        
        # Create embeddings for semantic search
        embeddings = []
        for doc in adgm_documents:
            embedding = self.embedding_model.encode(doc["content"])
            embeddings.append(embedding)
        
        embeddings = np.array(embeddings).astype('float32')
        
        # Create FAISS index for vector similarity search
        dimension = embeddings.shape[1]
        self.index = faiss.IndexFlatIP(dimension)  # Inner product for cosine similarity
        
        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(embeddings)
        self.index.add(embeddings)
    
    def retrieve_relevant_context(self, query: str, top_k: int = 5) -> List[Dict]:
        """Retrieve relevant ADGM legal context using semantic search"""
        query_embedding = self.embedding_model.encode([query])
        query_embedding = query_embedding.astype('float32')
        faiss.normalize_L2(query_embedding)
        
        scores, indices = self.index.search(query_embedding, top_k)
        
        relevant_docs = []
        for i, idx in enumerate(indices[0]):
            if scores[0][i] > 0.25:  # Similarity threshold
                relevant_docs.append({
                    "source": self.documents[idx]["source"],
                    "content": self.documents[idx]["content"],
                    "similarity": float(scores[0][i])
                })
        
        return relevant_docs

class DocumentProcessor:
    """Handle document processing for both .docx and .pdf files"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text content from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    @staticmethod
    def add_inline_comments(doc: Document, issues: List[Dict], doc_type: str) -> Document:
        """Add inline comments and highlights to document for flagged content"""
        try:
            processed_paragraphs = set()  # Track processed paragraphs
            
            for issue in issues:
                text_snippet = issue.get("text_snippet", "").strip()
                if not text_snippet or len(text_snippet) < 10:
                    continue
                
                # Find matching paragraphs
                search_text = text_snippet[:100].lower().strip()
                
                for para_idx, paragraph in enumerate(doc.paragraphs):
                    if para_idx in processed_paragraphs:
                        continue
                        
                    para_text = paragraph.text.lower().strip()
                    if len(para_text) < 10:
                        continue
                    
                    # Check if paragraph contains the flagged content
                    if DocumentProcessor._text_similarity(search_text, para_text) > 0.6:
                        # Highlight the paragraph
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        
                        # Add comment at the end of paragraph
                        comment_text = (
                            f"\n[ADGM COMPLIANCE REVIEW - {issue['severity'].upper()} PRIORITY]\n"
                            f"Issue: {issue['issue']}\n"
                            f"Legal Reference: {issue['rule_reference']}\n"
                            f"Recommendation: {issue['suggestion']}\n"
                            f"Confidence: {issue['confidence']:.2f}"
                        )
                        
                        # Create comment paragraph
                        comment_para = paragraph._element.getparent().insert(
                            paragraph._element.getparent().index(paragraph._element) + 1,
                            paragraph._element.__class__()
                        )
                        comment_para = paragraph.__class__(comment_para, paragraph._parent)
                        
                        comment_run = comment_para.add_run(comment_text)
                        comment_run.font.color.rgb = RGBColor(204, 0, 0)  # Red color
                        comment_run.font.size = Pt(9)
                        comment_run.bold = True
                        
                        # Add border/background effect
                        comment_para.paragraph_format.left_indent = Pt(36)
                        comment_para.paragraph_format.right_indent = Pt(36)
                        
                        processed_paragraphs.add(para_idx)
                        break
            
            # Add summary comment at the beginning
            if issues:
                summary_para = doc.paragraphs[0]._element.getparent().insert(0, doc.paragraphs[0]._element.__class__())
                summary_para = doc.paragraphs[0].__class__(summary_para, doc.paragraphs[0]._parent)
                
                summary_text = (
                    f"=== ADGM COMPLIANCE REVIEW SUMMARY ===\n"
                    f"Document Type: {doc_type}\n"
                    f"Review Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                    f"Issues Found: {len(issues)}\n"
                    f"High Priority: {len([i for i in issues if i['severity'] == 'High'])}\n"
                    f"Medium Priority: {len([i for i in issues if i['severity'] == 'Medium'])}\n"
                    f"Low Priority: {len([i for i in issues if i['severity'] == 'Low'])}\n"
                    f"Status: {'❌ NON-COMPLIANT' if any(i['severity'] == 'High' for i in issues) else '⚠️ REVIEW REQUIRED'}\n"
                    f"=======================================\n"
                )
                
                summary_run = summary_para.add_run(summary_text)
                summary_run.font.color.rgb = RGBColor(0, 0, 153)  # Blue color
                summary_run.font.size = Pt(10)
                summary_run.bold = True
                
        except Exception as e:
            print(f"Warning: Could not add inline comments: {e}")
        
        return doc
    
    @staticmethod
    def _text_similarity(text1: str, text2: str) -> float:
        """Calculate text similarity for paragraph matching"""
        words1 = set(text1.split())
        words2 = set(text2.split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = len(words1.intersection(words2))
        union = len(words1.union(words2))
        
        return intersection / union if union > 0 else 0.0

class ADGMCorporateAgent:
    def __init__(self):
        """Initialize ADGM Corporate Agent with RAG system"""
        self.rag_system = RAGSystem()
        self.groq_client = None
        self.doc_processor = DocumentProcessor()
        
        # Document type identification patterns
        self.doc_type_patterns = {
            "Articles of Association": [
                "articles of association", "articles", "company constitution",
                "memorandum and articles", "company rules", "constitutional document"
            ],
            "Memorandum of Association": [
                "memorandum of association", "memorandum", "company objects",
                "incorporation memorandum", "formation document"
            ],
            "Board Resolution": [
                "board resolution", "directors' resolution", "board meeting",
                "resolved that", "board of directors", "directors hereby resolve"
            ],
            "Board Resolution for Incorporation": [
                "board resolution", "incorporation resolution", "resolution for incorporation",
                "directors resolve to incorporate", "incorporation authorization"
            ],
            "Board Resolution for Licensing": [
                "board resolution", "licensing resolution", "resolution for licensing",
                "directors resolve to apply", "license application authorization"
            ],
            "UBO Declaration": [
                "ultimate beneficial owner", "ubo declaration", "beneficial ownership",
                "ownership structure", "controlling interest", "ubo form"
            ],
            "UBO Declaration Form": [
                "ultimate beneficial owner", "ubo declaration form", "beneficial ownership form",
                "ubo disclosure", "ownership declaration"
            ],
            "Employment Contract": [
                "employment contract", "employment agreement", "terms of employment",
                "employee", "employer", "salary", "compensation", "service agreement"
            ],
            "Employment Contract Template": [
                "employment contract template", "standard employment contract",
                "employment agreement template", "contract template"
            ],
            "Register of Members": [
                "register of members", "shareholder register", "member register",
                "share register", "membership list", "register of shareholders"
            ],
            "Register of Members and Directors": [
                "register of members and directors", "register of members", "register of directors",
                "member and director register", "company register"
            ],
            "Shareholder Resolution": [
                "shareholder resolution", "shareholders' resolution", "general meeting",
                "shareholders hereby resolve", "ordinary resolution", "special resolution"
            ],
            "Incorporation Application Form": [
                "incorporation application", "application form", "company registration form",
                "incorporation request", "registration application"
            ],
            "License Application Form": [
                "license application", "licensing application", "permit application",
                "business license form", "license request"
            ],
            "Business Plan": [
                "business plan", "business strategy", "operational plan",
                "company plan", "business proposal"
            ],
            "Financial Projections": [
                "financial projections", "financial forecast", "financial plan",
                "budget projections", "financial statements"
            ],
            "Employee Handbook": [
                "employee handbook", "staff handbook", "employee manual",
                "company handbook", "employee guide"
            ],
            "HR Policies Manual": [
                "hr policies", "human resources policies", "hr manual",
                "personnel policies", "hr procedures"
            ],
            "Professional Qualification Certificates": [
                "qualification certificates", "professional certificates", "credentials",
                "certifications", "professional qualifications"
            ],
            "Compliance Manual": [
                "compliance manual", "compliance procedures", "regulatory compliance",
                "compliance guide", "compliance policies"
            ],
            "Workplace Health & Safety Policy": [
                "health and safety", "workplace safety", "safety policy",
                "occupational health", "safety procedures"
            ],
            "Data Protection Policy": [
                "data protection", "privacy policy", "data privacy",
                "gdpr policy", "data security policy"
            ]
        }

        # Enhanced process signatures for automatic detection
        self.process_signatures = {
            "Company Incorporation": {
                "core_docs": ["Articles of Association", "Memorandum of Association"],
                "common_docs": ["Board Resolution", "Board Resolution for Incorporation", 
                               "UBO Declaration", "UBO Declaration Form", "Register of Members",
                               "Register of Members and Directors", "Incorporation Application Form"],
                "keywords": ["incorporation", "company formation", "register", "articles", "memorandum"]
            },
            "Licensing Application": {
                "core_docs": ["License Application Form", "Business Plan"],
                "common_docs": ["Financial Projections", "Board Resolution for Licensing",
                               "Professional Qualification Certificates", "Compliance Manual"],
                "keywords": ["license", "permit", "business plan", "application", "qualification"]
            },
            "Employment Setup": {
                "core_docs": ["Employment Contract", "Employment Contract Template"],
                "common_docs": ["Employee Handbook", "HR Policies Manual", 
                               "Workplace Health & Safety Policy", "Data Protection Policy"],
                "keywords": ["employment", "employee", "hr", "contract", "handbook"]
            }
        }

    def identify_document_type(self, content: str) -> str:
        """Identify document type using comprehensive content analysis"""
        content_lower = content.lower()
        
        # Score each document type
        type_scores = {}
        for doc_type, keywords in self.doc_type_patterns.items():
            score = 0
            for keyword in keywords:
                if keyword in content_lower:
                    # Weight longer keywords more heavily
                    score += len(keyword.split())
            type_scores[doc_type] = score
        
        # Return highest scoring type
        if type_scores and max(type_scores.values()) > 0:
            return max(type_scores.items(), key=lambda x: x[1])[0]
        
        return "Unknown Document Type"

    def detect_legal_process_automatically(self, uploaded_doc_types: List[str], all_content: str = "") -> Tuple[str, float]:
        """Automatically detect legal process based on uploaded document types and content"""
        
        # Score each process based on document matches
        uploaded_set = set(uploaded_doc_types)
        process_scores = {}
        
        for process, signature in self.process_signatures.items():
            score = 0
            
            # Check document type matches
            core_matches = len(set(signature["core_docs"]).intersection(uploaded_set))
            common_matches = len(set(signature["common_docs"]).intersection(uploaded_set))
            
            # Weight core documents more heavily
            doc_score = (core_matches * 5) + (common_matches * 2)
            
            # Check content keywords
            content_lower = all_content.lower()
            keyword_score = sum(1 for keyword in signature["keywords"] if keyword in content_lower)
            
            total_score = doc_score + keyword_score
            process_scores[process] = total_score
        
        # Find best match
        if process_scores:
            best_process = max(process_scores.items(), key=lambda x: x[1])
            max_possible_score = 25  # Rough estimate for normalization
            confidence = min(best_process[1] / max_possible_score, 1.0)
            
            # Only return if confidence is reasonable
            if best_process[1] >= 3:  # Minimum threshold
                return best_process[0], confidence
        
        return "Unknown Process", 0.0

    def get_enhanced_document_requirements(self, process_type: str) -> List[str]:
        """Comprehensive document requirements for all processes"""
        requirements = {
            "Company Incorporation": [
                "Articles of Association",
                "Memorandum of Association",
                "Board Resolution for Incorporation",
                "UBO Declaration Form",
                "Register of Members and Directors",
                "Incorporation Application Form"
            ],
            "Licensing Application": [
                "License Application Form",
                "Business Plan",
                "Financial Projections",
                "Board Resolution for Licensing",
                "Professional Qualification Certificates",
                "Compliance Manual"
            ],
            "Employment Setup": [
                "Employment Contract Template",
                "Employee Handbook",
                "HR Policies Manual",
                "Workplace Health & Safety Policy",
                "Data Protection Policy"
            ]
        }
        
        return requirements.get(process_type, [])

    def analyze_document_with_rag(self, document_content: str, doc_type: str) -> List[Dict]:
        """Analyze document using RAG-retrieved ADGM legal context with Llama 3"""
        
        if not self.groq_client:
            return [{
                "document": doc_type,
                "section_heading": "System Error",
                "text_snippet": "",
                "issue": "Groq API client not initialized",
                "severity": "High",
                "rule_reference": "System Error",
                "suggestion": "Please provide valid Groq API key",
                "confidence": 1.0
            }]
        
        # Retrieve relevant legal context from RAG system
        legal_context = self.rag_system.retrieve_relevant_context(
            f"{doc_type} ADGM compliance requirements legal issues red flags", 
            top_k=7
        )
        
        if not legal_context:
            return [{
                "document": doc_type,
                "section_heading": "RAG Analysis",
                "text_snippet": "Legal context retrieval failed",
                "issue": "Not enough information from ADGM legal database",
                "severity": "Medium",
                "rule_reference": "RAG System",
                "suggestion": "Manual review recommended with additional ADGM legal resources",
                "confidence": 0.5
            }]
        
        # Create comprehensive prompt for Groq API with retrieved context
        context_text = "\n".join([
            f"Legal Source: {doc['source']}\nContent: {doc['content']}\n---"
            for doc in legal_context
        ])
        
        # Truncate document content if too long
        doc_content = document_content[:3000] + "..." if len(document_content) > 3000 else document_content
        
        prompt = f"""You are an ADGM legal compliance expert. Analyze the document ONLY using the provided legal context.

RETRIEVED ADGM LEGAL CONTEXT:
{context_text}

DOCUMENT TYPE: {doc_type}
DOCUMENT CONTENT:
{doc_content}

ANALYSIS INSTRUCTIONS:
1. Compare document content against ADGM legal requirements from the context above
2. Identify specific compliance violations, missing clauses, or incorrect provisions
3. For each issue found, provide:
   - Exact section/clause from document that has the problem
   - Specific compliance issue description  
   - Legal reference from the provided context
   - Specific suggestion for correction
4. Classify severity: High (legal invalidity), Medium (compliance risk), Low (best practice)
5. Only cite laws/regulations from the provided context above
6. If context lacks sufficient information, respond with "Not enough information"

REQUIRED JSON OUTPUT FORMAT:
[
  {{
    "section_heading": "specific section name from document",
    "text_snippet": "exact problematic text from document (max 200 chars)",
    "issue": "specific compliance problem description",
    "severity": "High|Medium|Low",
    "rule_reference": "exact legal source citation from provided context",
    "suggestion": "specific actionable correction",
    "confidence": 0.85
  }}
]

Respond ONLY with valid JSON array. No additional text."""

        try:
            # Call Groq API with Llama 3
            response = self.groq_client.chat.completions.create(
                model="llama3-8b-8192",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=2500
            )
            
            response_text = response.choices[0].message.content.strip()
            
            # Handle insufficient information response
            if "not enough information" in response_text.lower():
                return [{
                    "document": doc_type,
                    "section_heading": "Analysis Limitation",
                    "text_snippet": "",
                    "issue": "Insufficient ADGM legal context for comprehensive analysis",
                    "severity": "Medium",
                    "rule_reference": "RAG System Limitation",
                    "suggestion": "Manual review recommended with complete ADGM legal documentation",
                    "confidence": 1.0
                }]
            
            # Parse JSON response
            try:
                # Clean response text
                response_text = response_text.strip()
                if response_text.startswith('```json'):
                    response_text = response_text[7:]
                if response_text.endswith('```'):
                    response_text = response_text[:-3]
                
                issues = json.loads(response_text)
                
                # Validate and enhance each issue
                validated_issues = []
                for issue in issues:
                    if isinstance(issue, dict) and "issue" in issue:
                        # Add document type and ensure all required fields
                        issue["document"] = doc_type
                        issue.setdefault("section_heading", "Document Analysis")
                        issue.setdefault("text_snippet", "")
                        issue.setdefault("confidence", 0.8)
                        validated_issues.append(issue)
                
                return validated_issues if validated_issues else self._create_fallback_analysis(doc_type, document_content)
                
            except json.JSONDecodeError as e:
                # Fallback analysis if JSON parsing fails
                return [{
                    "document": doc_type,
                    "section_heading": "Parser Error",
                    "text_snippet": response_text[:200],
                    "issue": f"LLM response parsing failed: {str(e)}",
                    "severity": "Medium",
                    "rule_reference": "System Error",
                    "suggestion": "Manual review recommended",
                    "confidence": 0.3
                }]
                
        except Exception as e:
            return [{
                "document": doc_type,
                "section_heading": "API Error",
                "text_snippet": "",
                "issue": f"Failed to analyze document with Groq API: {str(e)}",
                "severity": "High",
                "rule_reference": "System Error",
                "suggestion": "Check API configuration and network connectivity",
                "confidence": 1.0
            }]

    def _create_fallback_analysis(self, doc_type: str, content: str) -> List[Dict]:
        """Create basic rule-based analysis as fallback"""
        issues = []
        content_lower = content.lower()
        
        # Basic red flag detection
        if "uae federal" in content_lower or "dubai court" in content_lower:
            issues.append({
                "document": doc_type,
                "section_heading": "Jurisdiction Clause",
                "text_snippet": "References to UAE Federal or Dubai Courts found",
                "issue": "Incorrect jurisdiction - must specify ADGM Courts",
                "severity": "High",
                "rule_reference": "ADGM Companies Regulations 2020, Article 6",
                "suggestion": "Replace with 'ADGM Courts' jurisdiction",
                "confidence": 0.9
            })
        
        if not re.search(r"adgm|abu dhabi global market", content_lower):
            issues.append({
                "document": doc_type,
                "section_heading": "ADGM Reference",
                "text_snippet": "No clear ADGM jurisdiction reference found",
                "issue": "Missing explicit ADGM jurisdiction specification",
                "severity": "Medium",
                "rule_reference": "ADGM Companies Regulations 2020",
                "suggestion": "Include explicit ADGM jurisdiction clause",
                "confidence": 0.7
            })
        
        return issues

    def check_document_completeness(self, uploaded_docs: List[str], process_type: str) -> Dict:
        """Check document completeness using enhanced ADGM requirements"""
        
        # Get enhanced requirements
        required_docs = self.get_enhanced_document_requirements(process_type)
        
        # Enhance with RAG-retrieved requirements
        checklist_context = self.rag_system.retrieve_relevant_context(
            f"{process_type} required documents checklist ADGM incorporation mandatory",
            top_k=3
        )
        
        if checklist_context:
            for context in checklist_context:
                content = context["content"].lower()
                for doc_type in self.doc_type_patterns.keys():
                    if doc_type.lower() in content and doc_type not in required_docs:
                        required_docs.append(doc_type)
        
        # Calculate completeness
        uploaded_set = set(uploaded_docs)
        required_set = set(required_docs)
        missing = list(required_set - uploaded_set)
        
        return {
            "complete": len(missing) == 0,
            "missing": missing,
            "uploaded_count": len(uploaded_docs),
            "required_count": len(required_docs),
            "process_type": process_type
        }

    def process_documents(self, files, process_type: str, groq_api_key: str) -> Dict[str, Any]:
        """Main processing function with complete RAG integration and automatic process detection"""
        
        if not files:
            return {
                "error": "No files uploaded. Please upload at least one document (.docx or .pdf)",
                "timestamp": datetime.now().isoformat()
            }
        
        # Initialize Groq client
        self.groq_client = Groq(api_key=groq_api_key.strip())
        
        uploaded_docs = []
        all_issues = []
        processed_files = []
        all_content = ""  # For automatic process detection
        
        # Process each uploaded file
        for file in files:
            try:
                # Determine file type and extract content
                file_extension = os.path.splitext(file.name)[1].lower()
                content = ""
                
                if file_extension == ".docx":
                    # Read and parse .docx file
                    doc = Document(file.name)
                    content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    doc_type = self.identify_document_type(content)
                elif file_extension == ".pdf":
                    # Extract text from PDF file
                    content = self.doc_processor.extract_text_from_pdf(file.name)
                    doc_type = self.identify_document_type(content)
                else:
                    all_issues.append({
                        "document": file.name,
                        "section_heading": "File Type",
                        "text_snippet": "",
                        "issue": f"Unsupported file type: {file_extension}",
                        "severity": "High",
                        "rule_reference": "File Processing Error",
                        "suggestion": "Please upload only .docx or .pdf files",
                        "confidence": 1.0
                    })
                    continue
                
                if not content.strip():
                    all_issues.append({
                        "document": file.name,
                        "section_heading": "File Content",
                        "text_snippet": "",
                        "issue": "Document appears to be empty or contains no readable text",
                        "severity": "High",
                        "rule_reference": "Document Processing Error",
                        "suggestion": "Ensure document contains text content and is not corrupted",
                        "confidence": 1.0
                    })
                    continue
                
                # Accumulate content for automatic process detection
                all_content += " " + content
                
                # Identify document type using enhanced pattern matching
                uploaded_docs.append(doc_type)
                
                # Perform RAG-powered legal analysis
                issues = self.analyze_document_with_rag(content, doc_type)
                all_issues.extend(issues)
                
                # Add inline comments to the document (only for DOCX files)
                if file_extension == ".docx":
                    commented_doc = self.doc_processor.add_inline_comments(doc, issues, doc_type)
                    
                    # Save the reviewed document
                    output_filename = f"REVIEWED_{doc_type.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    temp_path = os.path.join(tempfile.gettempdir(), output_filename)
                    commented_doc.save(temp_path)
                    processed_files.append((temp_path, output_filename))
                else:
                    # For PDF files, we'll just create a text report
                    output_filename = f"REVIEWED_{doc_type.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                    temp_path = os.path.join(tempfile.gettempdir(), output_filename)
                    with open(temp_path, 'w', encoding='utf-8') as f:
                        f.write(f"ADGM Compliance Review for {doc_type}\n")
                        f.write("=" * 50 + "\n\n")
                        f.write(f"Review Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                        
                        if issues:
                            f.write("COMPLIANCE ISSUES FOUND:\n")
                            f.write("-" * 30 + "\n\n")
                            for i, issue in enumerate(issues, 1):
                                f.write(f"{i}. Issue: {issue['issue']}\n")
                                f.write(f"   Severity: {issue['severity']}\n")
                                f.write(f"   Legal Reference: {issue['rule_reference']}\n")
                                f.write(f"   Recommendation: {issue['suggestion']}\n")
                                f.write(f"   Confidence: {issue['confidence']:.2f}\n")
                                f.write(f"   Context: {issue['text_snippet'][:120]}{'...' if len(issue['text_snippet']) > 120 else ''}\n\n")
                        else:
                            f.write("No compliance issues detected based on available ADGM legal context.\n")
                    
                    processed_files.append((temp_path, output_filename))
                
            except Exception as e:
                all_issues.append({
                    "document": getattr(file, 'name', 'Unknown File'),
                    "section_heading": "File Processing Error",
                    "text_snippet": "",
                    "issue": f"Failed to process document: {str(e)}",
                    "severity": "High",
                    "rule_reference": "System Error",
                    "suggestion": "Ensure file is a valid .docx or .pdf document and try again",
                    "confidence": 1.0
                })
        
        # Automatic process detection if not manually specified
        detected_process = process_type
        detection_confidence = 1.0
        
        if process_type == "Auto-Detect" or not process_type:
            detected_process, detection_confidence = self.detect_legal_process_automatically(uploaded_docs, all_content)
        
        # Check document completeness against ADGM requirements
        completeness = self.check_document_completeness(uploaded_docs, detected_process)
        
        # Generate comprehensive compliance report
        report = {
            "process": detected_process,
            "process_detection_confidence": detection_confidence,
            "documents_uploaded": len(uploaded_docs),
            "uploaded_document_types": uploaded_docs,
            "required_documents": completeness["required_count"],
            "required_document_list": self.get_enhanced_document_requirements(detected_process),
            "missing_documents": completeness["missing"],
            "issues_found": all_issues,
            "timestamp": datetime.now().isoformat(),
            "report_version": "v2.0"
        }
        
        return {
            "report": report,
            "processed_files": processed_files,
            "success": True,
            "completeness_message": self._generate_completeness_message(completeness, detected_process)
        }
    
    def _generate_completeness_message(self, completeness: Dict, process_type: str) -> str:
        """Generate human-readable completeness message matching task requirements"""
        if completeness["complete"]:
            return f"✅ All required documents for {process_type} have been uploaded."
        else:
            # Format exactly as specified in task requirements
            missing_docs = ", ".join([f"'{doc}'" for doc in completeness["missing"]])
            
            # Use the exact format from task specification
            return (
                f"It appears that you're trying to complete {process_type} in ADGM. "
                f"Based on our reference list, you have uploaded {completeness['uploaded_count']} "
                f"out of {completeness['required_count']} required documents. "
                f"The missing document(s) appear to be: {missing_docs}."
            )

    def create_example_documents(self):
        """Create example documents for demo purposes"""
        try:
            # Create example Articles of Association
            example_doc = Document()
            
            # Add title
            title = example_doc.add_heading("ARTICLES OF ASSOCIATION", 0)
            
            # Add content with intentional compliance issues for demonstration
            example_doc.add_paragraph("EXAMPLE COMPANY LIMITED")
            example_doc.add_paragraph("A Private Company Limited by Shares")
            
            example_doc.add_heading("1. INTERPRETATION", level=1)
            example_doc.add_paragraph(
                'In these Articles, unless the context otherwise requires:'
            )
            
            example_doc.add_heading("2. SHARE CAPITAL", level=1)
            example_doc.add_paragraph(
                "The authorized share capital of the Company is AED 150,000 divided into 150,000 ordinary shares of AED 1.00 each."
            )
            
            example_doc.add_heading("3. JURISDICTION", level=1)
            # Intentional error for demonstration
            example_doc.add_paragraph(
                "Any disputes arising shall be subject to the jurisdiction of Dubai Courts."
            )
            
            example_doc.add_heading("4. DIRECTORS", level=1)
            example_doc.add_paragraph(
                "The Company shall have not less than one Director."
            )
            
            # Save example document
            example_path = os.path.join(tempfile.gettempdir(), "Example_Articles_of_Association_BEFORE_REVIEW.docx")
            example_doc.save(example_path)
            
            return example_path
            
        except Exception as e:
            print(f"Error creating example document: {e}")
            return None

# Initialize agent
agent = ADGMCorporateAgent()

def process_legal_documents(files, process_type, groq_api_key):
    """Gradio interface function with complete RAG + inline commenting + automatic detection"""
    
    # Use the API key from environment variables if not provided
    if not groq_api_key:
        groq_api_key = GROQ_API_KEY
    
    if not groq_api_key or not groq_api_key.strip():
        return "❌ Groq API key not found. Please set it in your .env file.", None, None
    
    if not files:
        return "❌ Please upload at least one document (.docx or .pdf)", None, None
    
    try:
        # Process documents with RAG analysis and automatic detection
        result = agent.process_documents(files, process_type, groq_api_key)
        
        if not result.get("success"):
            return f"❌ Processing Error: {result.get('error', 'Unknown error occurred')}", None, None
        
        report = result["report"]
        completeness_msg = result.get("completeness_message", "")
        
        # Generate comprehensive report display
        report_text = f"""
🏢 ADGM Corporate Agent - Enhanced RAG-Powered Document Review
============================================================
📋 Detected Process: {report['process']} (Confidence: {report['process_detection_confidence']:.2f})
📁 Documents Uploaded: {report['documents_uploaded']}
📄 Document Types: {', '.join(report['uploaded_document_types'])}
📋 Required Documents: {report['required_documents']}

{completeness_msg}

🚨 COMPLIANCE ANALYSIS ({len(report['issues_found'])} Issues Found):
"""
        
        if report['issues_found']:
            # Group issues by severity
            high_issues = [i for i in report['issues_found'] if i['severity'] == 'High']
            medium_issues = [i for i in report['issues_found'] if i['severity'] == 'Medium'] 
            low_issues = [i for i in report['issues_found'] if i['severity'] == 'Low']
            
            report_text += f"""
🔴 HIGH PRIORITY ({len(high_issues)}): {', '.join([i['issue'][:50]+'...' for i in high_issues[:3]])}
🟡 MEDIUM PRIORITY ({len(medium_issues)}): {', '.join([i['issue'][:50]+'...' for i in medium_issues[:3]])}
🟢 LOW PRIORITY ({len(low_issues)}): {', '.join([i['issue'][:50]+'...' for i in low_issues[:3]])}

DETAILED FINDINGS:
"""
            
            for i, issue in enumerate(report['issues_found'][:10], 1):
                severity_emoji = "🔴" if issue['severity'] == 'High' else "🟡" if issue['severity'] == 'Medium' else "🟢"
                report_text += f"""
{i}. {severity_emoji} [{issue['document']}] {issue['section_heading']}
   📍 Issue: {issue['issue']}
   ⚖️  Legal Basis: {issue['rule_reference']}
   💡 Recommendation: {issue['suggestion']}
   🎯 Confidence: {issue['confidence']:.2f}
   📝 Context: {issue['text_snippet'][:120]}{'...' if len(issue['text_snippet']) > 120 else ''}
"""
        else:
            report_text += "\n✅ No compliance issues detected based on available ADGM legal context.\n"
        
        # Add missing documents section
        if report['missing_documents']:
            report_text += f"""

📋 MISSING REQUIRED DOCUMENTS:
"""
            for missing_doc in report['missing_documents']:
                report_text += f"❌ {missing_doc}\n"
        
        report_text += f"""

📊 ANALYSIS SUMMARY:
• Process Detection: {'Automatic' if report['process_detection_confidence'] < 1.0 else 'Manual'} ({report['process_detection_confidence']:.2f} confidence)
• RAG System: Retrieved relevant ADGM legal context using FAISS vector database
• LLM Analysis: Powered by Llama 3 via Groq API
• Legal Sources: ADGM Companies Regulations 2020, Employment Regulations 2019, UBO Regulations 2021
• Document Processing: Inline comments and highlights added to reviewed documents
• Compliance Status: {'❌ NON-COMPLIANT' if any(i['severity'] == 'High' for i in report['issues_found']) else '⚠️ REVIEW RECOMMENDED' if report['issues_found'] else '✅ COMPLIANT'}

⏰ Analysis completed: {report['timestamp']}
🔄 Report Version: {report['report_version']}
"""
        
        # Prepare JSON report for download
        json_report_content = json.dumps(report, indent=2)
        json_temp_path = os.path.join(tempfile.gettempdir(), f"ADGM_Compliance_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
        with open(json_temp_path, 'w', encoding='utf-8') as f:
            f.write(json_report_content)
        
        # Prepare reviewed documents for download
        processed_files = result.get("processed_files", [])
        if processed_files:
            # Create zip file with all reviewed documents
            zip_temp_path = os.path.join(tempfile.gettempdir(), f"ADGM_Reviewed_Documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip")
            with zipfile.ZipFile(zip_temp_path, 'w') as zipf:
                for file_path, filename in processed_files:
                    if os.path.exists(file_path):
                        zipf.write(file_path, filename)
            
            return report_text, json_temp_path, zip_temp_path
        else:
            return report_text, json_temp_path, None
        
    except Exception as e:
        error_msg = f"❌ Critical Error: {str(e)}\n\nPlease check:\n1. Groq API key validity\n2. File format (.docx or .pdf)\n3. Network connectivity"
        return error_msg, None, None

def create_example_document():
    """Create and download example document for testing"""
    try:
        example_path = agent.create_example_documents()
        if example_path and os.path.exists(example_path):
            return example_path
        else:
            return None
    except Exception as e:
        print(f"Error creating example: {e}")
        return None

# Create comprehensive Gradio interface
with gr.Blocks(title="Enhanced ADGM Corporate Agent with Auto-Detection", theme=gr.themes.Soft()) as demo:
    
    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### 🔧 Configuration")
            
            process_type = gr.Dropdown(
                choices=[
                    "Auto-Detect",
                    "Company Incorporation",
                    "Licensing Application", 
                    "Employment Setup"
                ],
                value="Auto-Detect",
                label="📋 Legal Process Type"
            )
            
            gr.Markdown("*Select 'Auto-Detect' to automatically identify the process from your documents*")
            
            file_upload = gr.File(
                file_count="multiple",
                file_types=[".docx", ".pdf"],
                label="📄 Upload Legal Documents"
            )
            
            gr.Markdown("*Upload .docx or .pdf files for ADGM compliance review*")
            
            with gr.Row():
                submit_btn = gr.Button(
                    "🔍 Analyze Documents with Enhanced RAG",
                    variant="primary",
                    size="lg"
                )
                
                example_btn = gr.Button(
                    "📥 Download Example Document",
                    variant="secondary"
                )
        
        with gr.Column(scale=2):
            gr.Markdown("### 📊 Analysis Results")
            
            report_output = gr.Textbox(
                label="🔍 Enhanced RAG-Powered Compliance Analysis",
                lines=30,
                max_lines=40,
                show_copy_button=True
            )
    
    gr.Markdown("### 💾 Download Results")
    gr.Markdown("*Get structured compliance data and reviewed documents with inline comments*")
    
    with gr.Row():
        json_download = gr.File(
            label="📋 JSON Compliance Report",
            file_count="single"
        )
        
        docx_download = gr.File(
            label="📄 Reviewed Documents (.zip)",
            file_count="single"
        )
        
        example_download = gr.File(
            label="📥 Example Document (for testing)",
            file_count="single"
        )
    
    # Connect the processing function
    submit_btn.click(
        process_legal_documents,
        inputs=[file_upload, process_type, gr.State(GROQ_API_KEY)],
        outputs=[report_output, json_download, docx_download]
    )
    
    # Connect example document creation
    example_btn.click(
        create_example_document,
        outputs=[example_download]
    )

if __name__ == "__main__":
    demo.launch(share=True, server_name="0.0.0.0", server_port=7860)